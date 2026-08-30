import streamlit as st
from sentence_transformers import SentenceTransformer, util
import PyPDF2
import torch
import pandas as pd
import re
import sqlite3
from docx import Document
import io
from renderer import TemplateRenderer

st.set_page_config(page_title="Smart ATS & Candidate Portal", layout="wide")

# --- 1. DATABASE SETUP ---
def init_db():
    conn = sqlite3.connect('recruitment.db')
    c = conn.cursor()
    # Create Users Table (HR & Candidates)
    c.execute('''CREATE TABLE IF NOT EXISTS users 
                 (user_id TEXT PRIMARY KEY, name TEXT, password TEXT, role TEXT)''')
    # Create Certificates Table (Stores the Word Docs)
    c.execute('''CREATE TABLE IF NOT EXISTS certificates 
                 (cert_id INTEGER PRIMARY KEY AUTOINCREMENT, user_id TEXT, doc_type TEXT, file_data BLOB)''')
                 
    # Create Templates Table (Stores SVG and JSON for Document Engine)
    c.execute('''CREATE TABLE IF NOT EXISTS templates 
                 (template_id TEXT PRIMARY KEY, template_name TEXT, document_type TEXT, svg_data BLOB, template_json TEXT)''')
    
    # Create default HR Admin account
    c.execute("INSERT OR IGNORE INTO users (user_id, name, password, role) VALUES ('HR-001', 'GVK HR Team', 'admin123', 'HR')")
    conn.commit()
    conn.close()

init_db()

# --- 2. AUTHENTICATION SYSTEM ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['role'] = None
    st.session_state['user_id'] = None
    st.session_state['user_name'] = None

def login(user_id, password):
    conn = sqlite3.connect('recruitment.db')
    c = conn.cursor()
    c.execute("SELECT name, role FROM users WHERE user_id=? AND password=?", (user_id, password))
    user = c.fetchone()
    conn.close()
    
    if user:
        st.session_state['logged_in'] = True
        st.session_state['user_name'] = user[0]
        st.session_state['role'] = user[1]
        st.session_state['user_id'] = user_id
        st.rerun()
    else:
        st.error("❌ Invalid ID or Password")

def logout():
    st.session_state['logged_in'] = False
    st.session_state['role'] = None
    st.session_state['user_id'] = None
    st.session_state['user_name'] = None
    st.rerun()

# --- 3. AI & EXTRACTION ENGINES ---
@st.cache_resource
def load_models():
    return SentenceTransformer('all-MiniLM-L6-v2')
model = load_models()

def extract_text(uploaded_file):
    try:
        uploaded_file.seek(0) 
        reader = PyPDF2.PdfReader(uploaded_file)
        return " ".join([page.extract_text() for page in reader.pages if page.extract_text()]).strip()
    except:
        return ""

def classify_experience(text):
    text_lower = str(text).lower()
    if re.search(r'\d+\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience|professional experience', text_lower): return "Experienced"
    if re.search(r'fresher|entry[- ]level|internship', text_lower): return "Fresher"
    return "Unknown"

# --- 4. TEMPLATE ENGINE ---
def analyze_docx_template(file_bytes):
    file_bytes.seek(0) 
    doc = Document(file_bytes)
    full_text = " ".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " ".join([p.text for p in cell.paragraphs]) + " "

    doc_type = "Certificate" if "certificate" in full_text.lower() else "Document"
    placeholders = re.findall(r'\[(.*?)\]|\{(.*?)\}', full_text)
    clean_placeholders = {f"[{p[0]}]" if p[0] else f"{{{p[1]}}}" for p in placeholders}
    return doc_type, sorted(list(clean_placeholders))

def safe_replace(paragraph, placeholder, replacement):
    if placeholder not in paragraph.text: return
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, str(replacement))
            return
    paragraph.text = paragraph.text.replace(placeholder, str(replacement))

def fill_docx_template(file_bytes, inputs_dict):
    file_bytes.seek(0) 
    doc = Document(file_bytes)
    for p in doc.paragraphs:
        for ph, val in inputs_dict.items(): safe_replace(p, ph, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph, val in inputs_dict.items(): safe_replace(p, ph, val)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()

# ==========================================
#              USER INTERFACES
# ==========================================

if not st.session_state['logged_in']:
    st.title("🔐 Login Portal")
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.info("HR Default Login -> ID: **HR-001** | Pass: **admin123**")
        login_id = st.text_input("Enter ID (HR or Candidate ID)")
        login_pass = st.text_input("Enter Password", type="password")
        if st.button("Secure Login"):
            login(login_id, login_pass)

elif st.session_state['role'] == 'HR':
    st.sidebar.success(f"Logged in as: {st.session_state['user_name']} (HR)")
    st.sidebar.button("Logout", on_click=logout)
    
    tab1, tab2, tab3, tab4 = st.tabs(["1. ATS Scoring", "2. Register Candidates", "3. Word Certificates", "4. SVG Certificates (New)"])
    
    with tab1:
        st.header("Resume Scoring")
        job_description = st.text_area("Paste Job Requirements:")
        uploaded_files = st.file_uploader("Upload Resumes", type=["pdf"], accept_multiple_files=True)
        if st.button("Score Resumes"):
            if job_description and uploaded_files:
                resume_data, resume_texts = [], []
                for file in uploaded_files:
                    text = extract_text(file)
                    real_name = str(file.name).replace(".pdf", "").replace("_", " ")
                    resume_data.append({"Candidate Name": real_name, "Exp": classify_experience(text)})
                    resume_texts.append(text)
                
                if resume_texts:
                    job_embedding = model.encode(job_description, convert_to_tensor=True)
                    resume_embeddings = model.encode(resume_texts, convert_to_tensor=True)
                    cosine_scores = util.cos_sim(job_embedding, resume_embeddings)[0]
                    
                    results = []
                    for i, score in enumerate(cosine_scores):
                        results.append({"Name": resume_data[i]["Candidate Name"], "Score": f"{round(score.item() * 1.6 * 100, 1)}%"})
                    st.table(pd.DataFrame(results).sort_values(by="Score", ascending=False))
                    
    with tab2:
        st.header("Create Candidate Logins")
        st.write("Generate an ID and Password so the candidate can log in and download their certificate.")
        with st.form("register_candidate"):
            c_name = st.text_input("Candidate Name")
            c_id = st.text_input("Assign Candidate ID (e.g., GVK-1001)")
            c_pass = st.text_input("Assign Password")
            if st.form_submit_button("Create Candidate Account"):
                conn = sqlite3.connect('recruitment.db')
                try:
                    conn.execute("INSERT INTO users (user_id, name, password, role) VALUES (?, ?, ?, 'Candidate')", (c_id, c_name, c_pass))
                    conn.commit()
                    st.success(f"Account created! Give the candidate ID '{c_id}' and their password.")
                except sqlite3.IntegrityError:
                    st.error("This Candidate ID already exists!")
                conn.close()

    with tab3:
        st.header("Generate & Send to Candidate Portal")
        
        # Fetch available candidates from DB
        conn = sqlite3.connect('recruitment.db')
        candidates = pd.read_sql_query("SELECT user_id, name FROM users WHERE role='Candidate'", conn)
        conn.close()
        
        if candidates.empty:
            st.warning("No candidates registered yet. Go to Tab 2 to create a candidate ID.")
        else:
            selected_id = st.selectbox("Select Candidate to receive Certificate:", candidates['user_id'] + " - " + candidates['name'])
            actual_id = selected_id.split(" - ")[0]
            actual_name = selected_id.split(" - ")[1]
            
            docx_template = st.file_uploader("Upload Certificate Template (.docx)", type=['docx'])
            
            if docx_template:
                doc_type, placeholders = analyze_docx_template(docx_template)
                with st.form("cert_form"):
                    st.write(f"Generating for: **{actual_name}** ({actual_id})")
                    inputs = {}
                    for ph in placeholders:
                        if "name" in ph.lower():
                            st.info(f"✨ {ph} will be filled with: {actual_name}")
                            inputs[ph] = actual_name
                        else:
                            inputs[ph] = st.text_input(f"Value for {ph}:")
                            
                    if st.form_submit_button("Generate & Save to Candidate Portal"):
                        generated_doc = fill_docx_template(docx_template, inputs)
                        
                        # Save BLOB to SQLite
                        conn = sqlite3.connect('recruitment.db')
                        conn.execute("INSERT INTO certificates (user_id, doc_type, file_data) VALUES (?, ?, ?)", 
                                     (actual_id, f"{actual_name}_{doc_type}.docx", generated_doc))
                        conn.commit()
                        conn.close()
                        st.success(f"✅ Certificate securely saved! {actual_name} can now log in to download it.")
                        st.balloons()
                        
    with tab4:
        st.header("Generate AI-Ready SVG Certificates")
        st.info("Upload an SVG template with {{placeholders}}. The system will render a high-quality PDF using ReportLab.")
        
        conn = sqlite3.connect('recruitment.db')
        candidates_svg = pd.read_sql_query("SELECT user_id, name FROM users WHERE role='Candidate'", conn)
        conn.close()
        
        if candidates_svg.empty:
            st.warning("No candidates registered yet.")
        else:
            sel_id = st.selectbox("Select Candidate:", candidates_svg['user_id'] + " - " + candidates_svg['name'], key="svg_sel")
            act_id = sel_id.split(" - ")[0]
            act_name = sel_id.split(" - ")[1]
            
            svg_template = st.file_uploader("Upload SVG Template (.svg)", type=['svg'])
            
            if svg_template:
                svg_bytes = svg_template.read()
                svg_text = svg_bytes.decode('utf-8')
                
                # Automatically detect {{fields}}
                import re
                placeholders = list(set(re.findall(r'\{\{(.*?)\}\}', svg_text)))
                
                with st.form("svg_cert_form"):
                    st.write(f"Detected {len(placeholders)} fields in SVG.")
                    inputs = {}
                    for ph in placeholders:
                        if "name" in ph.lower():
                            st.info(f"✨ {ph} will be filled with: {act_name}")
                            inputs[ph] = act_name
                        else:
                            inputs[ph] = st.text_input(f"Value for {ph}:")
                            
                    if st.form_submit_button("Render PDF & Save to Portal"):
                        renderer = TemplateRenderer()
                        try:
                            # 1. Render the PDF
                            pdf_bytes = renderer.render_svg_to_pdf(svg_bytes, inputs)
                            
                            # 2. Save PDF BLOB to SQLite
                            conn = sqlite3.connect('recruitment.db')
                            conn.execute("INSERT INTO certificates (user_id, doc_type, file_data) VALUES (?, ?, ?)", 
                                         (act_id, f"{act_name}_Certificate.pdf", pdf_bytes))
                            
                            # 3. Save Template to database for Future AI Training (Path B)
                            import json
                            template_json = json.dumps({"fields": placeholders})
                            conn.execute("INSERT OR REPLACE INTO templates (template_id, template_name, document_type, svg_data, template_json) VALUES (?, ?, ?, ?, ?)",
                                         (svg_template.name, svg_template.name, "SVG_Certificate", svg_bytes, template_json))
                            conn.commit()
                            conn.close()
                            
                            st.success(f"✅ High-quality PDF generated and saved for {act_name}!")
                            st.balloons()
                        except Exception as e:
                            st.error(f"Error rendering SVG: {e}")

elif st.session_state['role'] == 'Candidate':
    st.sidebar.success(f"Logged in as: {st.session_state['user_name']}")
    st.sidebar.button("Logout", on_click=logout)
    
    st.title("🎓 My Candidate Portal")
    st.write(f"Welcome back, **{st.session_state['user_name']}**! Below are your official documents from GVK Industry.")
    
    conn = sqlite3.connect('recruitment.db')
    my_certs = pd.read_sql_query(f"SELECT cert_id, doc_type FROM certificates WHERE user_id='{st.session_state['user_id']}'", conn)
    
    if my_certs.empty:
        st.info("No certificates have been issued to your account yet.")
    else:
        for index, row in my_certs.iterrows():
            st.write("---")
            st.subheader(f"📄 {row['doc_type']}")
            
            # Retrieve the specific BLOB data for download
            c = conn.cursor()
            c.execute("SELECT file_data FROM certificates WHERE cert_id=?", (row['cert_id'],))
            file_data = c.fetchone()[0]
            
            mime_type = "application/pdf" if row['doc_type'].endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            st.download_button(
                label="⬇️ Download Document",
                data=file_data,
                file_name=row['doc_type'],
                mime=mime_type
            )
    conn.close()